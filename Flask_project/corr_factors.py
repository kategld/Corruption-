import nltk
import re
import string
from docx import Document
from docx.enum.text import WD_COLOR_INDEX 
from tensorflow.keras.preprocessing.sequence import pad_sequences
import numpy as np
import pandas as pd
nltk.download("stopwords")
from nltk.corpus import stopwords
from pymystem3 import Mystem
mystem = Mystem()

russian_stopwords = stopwords.words("russian")

class_dict={
    "0":"Широта дискреционных полномоция",
    "1":"Определение компетенции по формуле 'вправе'",
    "2":"Выборочное изменение прав и обязанностей",
    "3":"Чрезмерная свобода подзаконного норматворчества",
    "4":"Принятие нормативного правового акта за пределами компетенции",
    "7":"Заполнение законодательных пробелов",
    "5":"Отсутствие или неполнота административных процедур",
    "9":"Отказ от конкурсных аукционных процедур",
    "8":"Наличие завышенных требований",
    "6":"Злоупотребление правом заявителя",
    "10":"Юридико-лингвистическая неопределенность"
}

from natasha import (
    Segmenter,
    MorphVocab,
    NewsEmbedding,
    NewsMorphTagger,
    NewsSyntaxParser,
    NewsNERTagger,
    PER,
    NamesExtractor,
    DatesExtractor,
    AddrExtractor,
    Doc
)

segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
morph_tagger = NewsMorphTagger(emb)
syntax_parser = NewsSyntaxParser(emb)

def find_paper(text):
    paper_text=""
    if("'") in text:
        paper_text="'".join(text.split("'")[:2])+"'"
    elif('"') in text:     
        paper_text='"'.join(text.split('"')[:2])+'"'
    elif("«" and "»") in text:
         paper_text=text.split("»")[0]+"»"
    return paper_text

def find_ref(txt):
    index_words=["закон","постановление","приказ","указ","распоряжение","положение"]
    doc = Doc(txt)
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    for token in doc.tokens:
        token.lemmatize(morph_vocab)
    ind_dict=[]
    try:
        sent=doc.sents[0]
        for token in sent.tokens:
             if (token.lemma in index_words):
                ind_dict.append((token.start-1,token.stop))
    except:
        pass
    papers=[]
    try:
        for i in range(len(ind_dict)-1):
            text=sent.text[ind_dict[i][0]:ind_dict[i+1][0]]
            text=find_paper(text)
            if text:
                papers.append({(ind_dict[i][0],ind_dict[i][0]+len(text)):text})
        text=sent.text[ind_dict[-1][0]:]
        text=find_paper(text)
        if text:
            papers.append({(ind_dict[-1][0],ind_dict[-1][0]+len(text)):text})
    except:
        pass
    return papers    

def preprocess(text):
    text = text.lower() 
    text=text.strip()  
    text=re.compile('<*>\n').sub('', text) 
    text = re.compile('[%s]' % re.escape(string.punctuation)).sub(' ', text)  
    text = re.sub('\s+', ' ', text)  
    text = re.sub(r'\[[0-9]*\]',' ',text) 
    text=re.sub(r'[^\w\s]', '', str(text).lower().strip())
    text = re.sub(r'\d',' ',text) 
    text = re.sub(r'\s+',' ',text) 
    return text
def lemmatizer(string,stopwords):
    tokens = mystem.lemmatize(string)
    tokens = [token for token in tokens if token not in stopwords]
    text = " ".join(tokens)
    text=text.replace('\n',"")
    return text
def finalpreprocess(string,stopwords):
    return lemmatizer(preprocess(string),stopwords)

def find_classes(l_classes):
    positions=[]
    for i in range(len(l_classes)):
        if l_classes[i]==1:
            positions.append(str(i))
    if(not positions):
        return np.nan
    else:
        return " ".join(positions)

def change_doc(doc_name,tokenizer,load_model):
    document = Document(doc_name)
    text_puncts=[]
    for p in document.paragraphs:
        text_puncts.append(p.text)
        if find_ref(p.text):
            for run in p.runs:
                font = run.font
                font.highlight_color=WD_COLOR_INDEX.GRAY_25
            comment = p.add_comment('Возможно наличие фактора "Нормативные коллизии" с внешними документами',author='SEQUORA',initials= 'sqr') 
    sentences=[]
    for txt in text_puncts:
        txt=preprocess(txt)
        doc = Doc(txt)
        doc.segment(segmenter)
        doc.tag_morph(morph_tagger)
        for token in doc.tokens:
            token.lemmatize(morph_vocab)
        tokens=[t.lemma for t in doc.tokens]
        tokens = [token for token in tokens if token not in russian_stopwords]
        #if (tokens):
        sentences.append(tokens)
    test_x = np.array(sentences)
    test_sequences = tokenizer.texts_to_sequences(test_x)
    max_len=120
    padding_type='post'
    trunc_type='post'
    Xtest = pad_sequences(test_sequences,  maxlen=max_len, padding=padding_type, truncating=trunc_type)
    y_predicted = load_model.predict(Xtest)
    y_predicted=np.where(y_predicted > 0.5, 1, 0)
    cl=[find_classes(i) for i in y_predicted]
    class_data=pd.DataFrame(
        {"text":text_puncts,
        "class":cl})
    class_data=class_data.dropna()
    class_data=class_data.replace({"class":class_dict})
    corr_nums=list(class_data.index)
    for p in range(len(document.paragraphs)):
        if p in corr_nums:
            for run in document.paragraphs[p].runs:
                font = run.font
                font.highlight_color=WD_COLOR_INDEX.GRAY_25
            comment = document.paragraphs[p].add_comment("Возможно наличие фактора "+'"'+class_data.loc[p]["class"]+'"',author='SEQUORA',initials= 'sqr') 
    return document


